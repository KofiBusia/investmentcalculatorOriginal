# Generate lecture script Word document for both Excel valuation models.

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Colour palette ────────────────────────────────────────────
NAVY  = RGBColor(0x0D, 0x2B, 0x55)   # dark navy heading
GOLD  = RGBColor(0xC9, 0xA0, 0x2C)   # gold accent
TEAL  = RGBColor(0x09, 0x7A, 0x6E)   # teal for model 2
GRAY  = RGBColor(0x4A, 0x4A, 0x4A)   # body text
LGRAY = RGBColor(0x88, 0x88, 0x88)   # light gray notes

# ── Helper functions ──────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()

def heading1(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = color
    return p

def heading2(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = color
    return p

def heading3(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    return p

def instructor_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.35)
    run = p.add_run(f"[INSTRUCTOR NOTE: {text}]")
    run.italic    = True
    run.font.size = Pt(10)
    run.font.color.rgb = LGRAY
    return p

def formula_box(doc, formula_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run(formula_text)
    run.bold      = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = NAVY
    # shade the paragraph
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  'EEF4FB')
    p._p.get_or_add_pPr().append(shading)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 72)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.4 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    return p

def step(doc, number, heading_text, body_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"Step {number}:  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = GOLD
    r2 = p.add_run(heading_text)
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = NAVY
    body(doc, body_text, indent=True)

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("EQUITY VALUATION MODELS")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Full Lecturer Script  ·  Step-by-Step Guide")
r.font.size = Pt(14)
r.font.color.rgb = GOLD

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\nModel 1: DCF — Free Cash Flow to Equity (FCFE)\nModel 2: Residual Income Valuation Model (RIM)")
r.font.size = Pt(13)
r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("InvestIQ  ·  Young Investors Acceleration Programme (YIAP)")
r.font.size = Pt(11)
r.font.color.rgb = LGRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Instructor: Kofi Kyei")
r.font.size = Pt(11)
r.font.color.rgb = LGRAY

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# HOW TO USE THIS DOCUMENT
# ══════════════════════════════════════════════════════════════
heading1(doc, "How to Use This Script")
body(doc, "This document contains the full wording you will say to your students as you walk them through both Excel valuation models. The script is written exactly as you would speak it in the classroom — conversational, clear, and progressive.")
body(doc, "Each section follows the same structure:")
bullet(doc, "HEADING — the topic being introduced")
bullet(doc, "What to say — the exact words, written in plain conversational English")
bullet(doc, "[INSTRUCTOR NOTE] — a private cue to yourself (actions, pauses, what to show on screen)")
body(doc, "Work through Model 1 first (DCF FCFE), then Model 2 (Residual Income). Allow approximately 45–60 minutes per model.")
divider(doc)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  MODEL 1 — DCF FCFE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "MODEL 1 — DCF Valuation Using Free Cash Flow to Equity (FCFE)")

divider(doc)
heading2(doc, "Opening — Setting the Stage")
body(doc, "Good morning everyone. Today we are going to value a real company, step by step, the way analysts at Goldman Sachs, JP Morgan, and every major investment bank in the world do it. We are not doing theory. We are doing a live valuation. By the end of this session, you will be able to open this Excel file, put in the numbers of any company, and arrive at a price — a fair value — for that company's shares.")
body(doc, "The model we are using is called the Discounted Cash Flow model, specifically using Free Cash Flow to Equity — or FCFE for short. This is one of the two main valuation models taught in the CFA Institute curriculum and it is used on every major stock exchange in the world.")
instructor_note(doc, "Open the Excel file: 'DCF VALUATION Model FCFE.xlsx'. Maximise the window so students can see it clearly.")

divider(doc)
heading2(doc, "SHEET 1 — Notes (The Foundation)")
heading3(doc, "What is FCFE and why does it matter?")
body(doc, "The first sheet is called Notes. Before we touch any numbers, I want you to understand the concept. Read with me.")
body(doc, "Free Cash Flow to Equity — FCFE — is the cash that is left over for the equity shareholders of a company AFTER the company has paid all its operating expenses, paid its taxes, reinvested in the business by buying equipment and assets, managed its working capital, and serviced its debt.")
body(doc, "In simple terms — if a company earned money today and took care of every obligation it has, the cash that remains belongs to you, the shareholder. That is FCFE.")
body(doc, "Now here is the key question we are answering in this model: If the company will generate this much cash for shareholders every year into the future, what is that stream of cash worth today? That answer is the intrinsic value of the company.")
body(doc, "The formula for FCFE is:")
formula_box(doc, "FCFE  =  Net Income  +  Non-Cash Charges  −  Capital Expenditure  −  Change in Working Capital  +  Net Borrowings")
body(doc, "Let me break this down for you one line at a time.")
bullet(doc, "Net Income — this is the profit the company made after tax. It is the starting point.")
bullet(doc, "Plus Non-Cash Charges — the biggest one here is Depreciation and Amortisation. Depreciation is not a real cash outflow — it is just an accounting entry. So we add it back because we want real cash, not accounting profit.")
bullet(doc, "Minus Capital Expenditure — this is the money the company spent buying property, plant, and equipment. This is real cash going out.")
bullet(doc, "Minus Change in Working Capital — if the business is growing, it needs more cash tied up in inventory and receivables. This is cash that is consumed by operations.")
bullet(doc, "Plus Net Borrowings — if the company borrowed more money than it repaid this year, that extra cash is also available to equity holders.")
body(doc, "Take a moment. Does everyone understand what we are building? We are calculating the real cash the company generates for its shareholders each year.")
instructor_note(doc, "Pause. Ask students: 'Why do we add back depreciation?' Wait for answers before continuing.")

divider(doc)
heading2(doc, "SHEET 2 — Dashboard (The Big Picture)")
heading3(doc, "What does the Dashboard show?")
instructor_note(doc, "Click on the 'Dashboard' tab.")
body(doc, "Good. Now click on the Dashboard tab. This is the summary page — the first thing a CEO or investor sees. It shows the intrinsic value per share, the current market price, and whether the stock is undervalued or overvalued.")
body(doc, "Think of the Dashboard as the answer page. Everything we are about to fill in feeds into this page. By the time we finish, this page will tell us: should you buy this stock, sell it, or hold it?")
body(doc, "You will see three key outputs here:")
bullet(doc, "Intrinsic Value Per Share — what the model says the stock is truly worth.")
bullet(doc, "Current Market Price — what the stock is trading at on the stock exchange right now.")
bullet(doc, "Margin of Safety — the gap between the two. If intrinsic value is higher than market price, the stock is potentially undervalued — and that is a buying opportunity.")
body(doc, "Do not change anything on this sheet. The numbers here are all automatically calculated from the inputs you will enter in the next sheets.")
instructor_note(doc, "Scroll slowly across the Dashboard so students can read the key cells. Do not linger — move on to Inputs.")

divider(doc)
heading2(doc, "SHEET 3 — Inputs (Where You Do The Work)")
heading3(doc, "This is the most important sheet in the model")
instructor_note(doc, "Click on the 'Inputs' tab. Zoom in so students can see the cells clearly.")
body(doc, "This is the engine room. Everything you type here drives the entire model. Let us go through each input carefully.")

step(doc, 1, "Company Name and Ticker",
     "At the top, type the name of the company you are valuing and its stock ticker. For example, if we are valuing Cocoa Foods Ghana Limited, type that. This is just a label — it does not affect the calculation.")

step(doc, 2, "Number of Historical Years",
     "This is how many years of historical data you have for the company. Typically you will use three to five years. We recommend five years because more data gives you a more reliable picture of how the company has been generating cash. If you only have three years of data, type 3. The model will adjust automatically.")

step(doc, 3, "Historical Financial Data",
     "Now we start entering the real numbers from the company's financial statements. For each year, you will need five numbers. I want you to write these down:")
bullet(doc, "Net Income — from the Income Statement (bottom line, after tax)")
bullet(doc, "Depreciation & Amortisation (Non-Cash Charges) — from the Cash Flow Statement, Operating Activities section")
bullet(doc, "Capital Expenditure — from the Cash Flow Statement, Investing Activities section (usually shown as a negative number — enter it as a positive)")
bullet(doc, "Change in Working Capital — from the Cash Flow Statement (change in current assets minus current liabilities, excluding cash)")
bullet(doc, "Net Borrowings — from the Cash Flow Statement, Financing Activities section (new debt raised minus debt repaid)")
body(doc, "Enter the data for each year in the yellow-highlighted cells. The model will compute the FCFE for each year automatically.")
instructor_note(doc, "Walk through one year of sample data live. Use round numbers for demonstration, e.g. Net Income = GHS 50m, D&A = GHS 8m, Capex = GHS 12m, ΔNWC = GHS 3m, Net Borrowings = GHS 5m. Show the computed FCFE = 50 + 8 − 12 − 3 + 5 = GHS 48m.")

step(doc, 4, "Cost of Equity",
     "This is the return that shareholders require to invest in this company. It answers the question: what annual return do I need to be compensated for the risk of owning this stock? We calculate it using the Capital Asset Pricing Model — CAPM. The formula is:")
formula_box(doc, "Cost of Equity  =  Risk-Free Rate  +  Beta  ×  (Market Return  −  Risk-Free Rate)")
body(doc, "Let me explain each component.", indent=True)
bullet(doc, "Risk-Free Rate — this is the return you would earn with zero risk. In Ghana, we use the 91-day Treasury bill rate. As of today you can get this from the Bank of Ghana website. In the US, analysts use the 10-year US Treasury yield.", 1)
bullet(doc, "Beta — this measures how volatile the stock is compared to the overall market. A beta of 1.0 means the stock moves exactly with the market. A beta above 1.0 means it is more volatile — higher risk, higher potential return. You find Beta on Bloomberg, Reuters, or you can calculate it from historical share price data.", 1)
bullet(doc, "Market Return — this is the historical average annual return of the stock market. For Ghana Stock Exchange, this is approximately 15 to 20 percent. For the S&P 500, analysts typically use 10 to 12 percent.", 1)
body(doc, "Enter these three numbers and the model will calculate your Cost of Equity automatically.", indent=True)

step(doc, 5, "Perpetual Growth Rate",
     "This is the rate at which you expect the company's cash flow to grow forever into the future. This is also called the terminal growth rate. A very important rule: this rate must always be LOWER than the Cost of Equity. If growth exceeds the discount rate, the formula breaks down mathematically. In practice, use a rate between 2 and 6 percent. For a Ghanaian company, you might use 4 to 5 percent. For a mature US company, analysts typically use 2 to 3 percent — roughly in line with long-term GDP growth.")

step(doc, 6, "Current Market Price and Shares Outstanding",
     "Enter the current share price from the stock exchange and the total number of shares the company has issued. These two numbers let the model compare its calculated intrinsic value to the market price — and tell you if the stock is cheap or expensive.")

instructor_note(doc, "After all inputs are entered, scroll back to the Dashboard and show how the numbers have populated. This is a powerful visual moment — let students see the model 'light up'.")

divider(doc)
heading2(doc, "SHEET 4 — FCFE Calculation (Seeing the Formula Work)")
heading3(doc, "How the model computes FCFE year by year")
instructor_note(doc, "Click on the 'FCFE Calculation' tab.")
body(doc, "This sheet shows you the workings. You can see the FCFE being calculated for each historical year using the formula we discussed. The model takes your inputs and applies:")
formula_box(doc, "FCFE  =  Net Income  +  D&A  −  Capex  −  ΔNWC  +  Net Borrowings")
body(doc, "Look at the bottom of this table. You will see the most recent year's FCFE — the last number in the column. This is the base number the model uses to project future cash flows. The model does NOT take an average. It uses the most recent year because that is the most current picture of what the company is generating.")
body(doc, "The model then projects this FCFE forward using your perpetual growth rate. For each future year:")
formula_box(doc, "Projected FCFE (Year t)  =  Last FCFE  ×  (1  +  g)^t")
body(doc, "So if the last FCFE was GHS 48 million and your growth rate is 4%, then Year 1 projected FCFE = 48 × 1.04 = GHS 49.9 million, Year 2 = 48 × 1.04² = GHS 51.9 million, and so on.")
instructor_note(doc, "Walk students through two rows of the projection table so they can see the formula in action.")

divider(doc)
heading2(doc, "SHEET 5 — DCF Valuation (Arriving at the Price)")
heading3(doc, "Discounting cash flows back to today")
instructor_note(doc, "Click on the 'DCF Valuation' tab.")
body(doc, "Now we arrive at the heart of the valuation. This sheet takes all those projected future cash flows and asks: what is each of these future cash flows worth in today's money?")
body(doc, "This is the concept of Present Value — or the Time Value of Money. A Ghana cedi today is worth more than a Ghana cedi in five years' time. Why? Because if you have money today, you can invest it and earn a return. So we discount future cash flows back using the Cost of Equity as our discount rate.")
formula_box(doc, "Present Value of FCFE (Year t)  =  FCFE_t  ÷  (1  +  Cost of Equity)^t")
body(doc, "The model does this for each projected year and sums them all up. That gives us the Present Value of the forecast period cash flows.")
body(doc, "But wait — a company does not stop operating after five years. So we need to capture the value of ALL cash flows beyond our forecast period. That is what the Terminal Value does.")
heading3(doc, "Understanding Terminal Value")
body(doc, "Terminal Value captures the value of the company from Year N+1 to infinity. The formula is:")
formula_box(doc, "Terminal Value  =  FCFE_(N+1)  ÷  (Cost of Equity  −  Perpetual Growth Rate)")
body(doc, "Then we discount that Terminal Value back to today:")
formula_box(doc, "PV of Terminal Value  =  Terminal Value  ÷  (1  +  Cost of Equity)^N")
body(doc, "The Total Equity Value of the company is:")
formula_box(doc, "Total Equity Value  =  PV of Forecast Cash Flows  +  PV of Terminal Value")
body(doc, "And the Intrinsic Value Per Share is simply:")
formula_box(doc, "Intrinsic Value Per Share  =  Total Equity Value  ÷  Shares Outstanding")
body(doc, "This is the number you compare to the current market price. If the intrinsic value is HIGHER than the market price, the stock may be undervalued. If it is LOWER, the stock may be overvalued.")
instructor_note(doc, "Show the final intrinsic value on screen. Compare it to the market price entered in the Inputs sheet. Ask students: 'Based on this model, would you buy this stock?'")

divider(doc)
heading2(doc, "SHEET 6 — Sensitivity Analysis (Testing Your Assumptions)")
heading3(doc, "What happens when our assumptions change?")
instructor_note(doc, "Click on the 'Sensitivity' tab.")
body(doc, "This is one of the most important sheets in the model — and unfortunately, one that many beginners skip. Do not skip it.")
body(doc, "Every number we entered is an assumption. The Cost of Equity is an assumption. The Growth Rate is an assumption. What if we are wrong? What if the growth rate is 3% instead of 5%? What if the Cost of Equity is 18% instead of 15%? The Sensitivity table shows you how the intrinsic value changes when you vary these two key assumptions.")
body(doc, "The rows represent different Cost of Equity values. The columns represent different Perpetual Growth Rates. Each cell in the table shows you the intrinsic value per share for that combination.")
body(doc, "Here is how to read it: look at the range of values across the table. If the intrinsic value is above the market price in MOST cells, the stock looks undervalued under most reasonable assumptions. If it is only above the market price in the very optimistic corner of the table, be cautious — the valuation is very sensitive to your assumptions.")
body(doc, "Professionals call this a 'bull case / base case / bear case' analysis. You always present a range of values, not a single number. No analyst will ever tell you a stock is worth exactly GHS 12.50. They will say it is worth between GHS 10.00 and GHS 15.00 under different assumptions.")
instructor_note(doc, "Point to the green-shaded cells (where intrinsic value > market price) and the red-shaded cells (where intrinsic value < market price). Ask: 'Under how many scenarios is this stock a buy?'")

divider(doc)
heading2(doc, "Closing — Model 1 Summary")
body(doc, "Excellent work. Let us recap what we just did.")
bullet(doc, "We understood what FCFE is — the real cash available to shareholders.")
bullet(doc, "We entered historical financial data from the company's statements.")
bullet(doc, "We set the Cost of Equity using CAPM.")
bullet(doc, "The model projected future cash flows and discounted them to today.")
bullet(doc, "Terminal Value captured cash flows beyond our forecast window.")
bullet(doc, "We compared intrinsic value to market price.")
bullet(doc, "We stress-tested our assumptions using sensitivity analysis.")
body(doc, "This is exactly how a buy-side equity analyst at a fund manager values a stock before recommending a buy or sell. You have just done a professional-grade valuation.")
instructor_note(doc, "Take a 10-minute break before starting Model 2.")
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  MODEL 2 — RESIDUAL INCOME VALUATION
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "MODEL 2 — Residual Income Valuation Model (RIM)", color=TEAL)

divider(doc)
heading2(doc, "Opening — Why a Different Model?", color=TEAL)
body(doc, "Welcome back. Now we move to the second model — the Residual Income Valuation Model, or RIM. A question you should be asking is: if DCF FCFE works, why do we need another model?")
body(doc, "Great question. The DCF FCFE model works well for companies that generate clear, measurable free cash flow — manufacturing companies, consumer goods companies, retailers. But banks and financial institutions are different. Their business is money itself. Their assets and liabilities are financial instruments. It is very difficult to isolate 'free cash flow' for a bank the way we can for a normal company.")
body(doc, "That is why, for banks and financial firms, analysts prefer the Residual Income Model. Instead of starting from cash flow, this model starts from the book value of the company — the net assets on the balance sheet — and then adds the extra value the company creates above and beyond what shareholders require.")
body(doc, "This model is used by Morgan Stanley, Deutsche Bank, and every major bank equity research team globally. It is also a core part of the CFA Institute Level 2 and Level 3 curriculum.")
instructor_note(doc, "Open the Excel file: 'RESIDUAL INCOME VALUATION MODEL.xlsx'. Maximise the window.")

divider(doc)
heading2(doc, "SHEET 1 — Cover (Understanding the Concept)", color=TEAL)
heading3(doc, "What is Residual Income?", color=TEAL)
body(doc, "The first sheet is the Cover — and it contains the most important concept in this model.")
body(doc, "Residual Income is the income a company earns ABOVE the minimum return that shareholders require. Think of it this way: if you invest in a company and you require a 15% annual return, but the company only earns 12% on your money — it has not met your expectations. That shortfall is a negative residual income.")
body(doc, "On the other hand, if the company earns 20% on your equity when you only required 15%, it has created VALUE above and beyond your expectations. That excess is positive residual income, and it adds to the company's worth.")
body(doc, "The formula is:")
formula_box(doc, "Residual Income  =  EPS  −  (Cost of Equity  ×  Book Value Per Share at start of year)")
body(doc, "Or expressed as a rate:")
formula_box(doc, "Residual Income  =  (ROE  −  Cost of Equity)  ×  Book Value Per Share")
body(doc, "Where ROE is Return on Equity — how much profit the company makes for every unit of shareholder equity. If ROE exceeds the Cost of Equity, the company is creating value. If ROE is below the Cost of Equity, the company is destroying value even if it is showing an accounting profit.")
instructor_note(doc, "Pause and ask students: 'A company shows a profit of GHS 10 million. Its equity is GHS 100 million and shareholders require 15% return. Is this company creating or destroying value?' Answer: 10m ÷ 100m = 10% ROE < 15% required. It is destroying value despite being profitable.")

divider(doc)
heading2(doc, "SHEET 2 — Instructions (Read Before You Touch Anything)", color=TEAL)
instructor_note(doc, "Click on the 'Instructions' tab.")
body(doc, "This sheet tells you exactly which cells to fill in and which to leave alone. The yellow cells are your input cells. The white cells with formulas should never be touched — they are the calculations.")
body(doc, "This is a professional financial model. It has been built so that you only need to enter data in clearly marked areas. If you type over a formula by mistake, the model breaks. Always work only in the yellow cells.")
body(doc, "The instructions also tell you which financial statements to use to find each input. We will go through this in detail now.")
instructor_note(doc, "Read the instructions sheet aloud with students, slowly. This sets expectations and prevents errors later.")

divider(doc)
heading2(doc, "SHEET 3 — Inputs (Entering the Company Data)", color=TEAL)
heading3(doc, "What you need before you start", color=TEAL)
instructor_note(doc, "Click on the 'Inputs' tab.")
body(doc, "Before we enter any data, let me tell you what documents you need to have open alongside this model:")
bullet(doc, "The company's Annual Report — specifically the last 3 to 5 years")
bullet(doc, "Income Statement — for Net Profit (Net Income) each year")
bullet(doc, "Balance Sheet — for Equity (Book Value of Equity) each year")
bullet(doc, "Any available EPS (Earnings Per Share) data — from the annual report or stock exchange")
instructor_note(doc, "Give students 2 minutes to locate these numbers from a sample annual report you have prepared.")

step(doc, 1, "Company Information",
     "At the top, enter the company name, the sector or industry, and the current date. This is for identification only — it does not affect the calculation.")

step(doc, 2, "CAPM Inputs — Cost of Equity",
     "Just as in the DCF model, we calculate Cost of Equity using CAPM.")
formula_box(doc, "Cost of Equity  =  Risk-Free Rate  +  Beta  ×  (Market Return  −  Risk-Free Rate)")
body(doc, "Enter the same three components you used before:", indent=True)
bullet(doc, "Risk-Free Rate: e.g. 91-Day T-Bill rate from Bank of Ghana", 1)
bullet(doc, "Beta: from Bloomberg, Reuters, or calculate from price history", 1)
bullet(doc, "Market Return: historical average annual return of the market index", 1)
body(doc, "The model calculates Cost of Equity automatically from these inputs.", indent=True)

step(doc, 3, "Book Value of Equity",
     "This is the current total shareholders' equity from the Balance Sheet — the net assets of the company as recorded in the accounts. You will find this at the bottom of the Balance Sheet under 'Total Equity' or 'Shareholders' Funds'. Enter the figure for the most recent financial year.")

step(doc, 4, "Earnings Per Share (EPS) Forecasts",
     "For each year of your forecast period — usually 3 to 5 years — enter your expected EPS. You have two options here:")
bullet(doc, "Use analyst consensus forecasts from Bloomberg or Reuters if available", 1)
bullet(doc, "Calculate your own forecasts based on the company's historical ROE growth trend", 1)
body(doc, "If you are forecasting yourself, a simple method is: look at the average EPS growth over the last 3 years and apply that growth rate to the most recent EPS.", indent=True)
instructor_note(doc, "Demonstrate: if EPS for Year -3, -2, -1 were GHS 0.80, GHS 0.90, GHS 1.00 — average growth is 11.8%. Project Year 1 EPS = 1.00 × 1.118 = GHS 1.12, and so on.")

step(doc, 5, "Dividend Policy",
     "This model allows three dividend policy options. Select the one that matches the company:")
bullet(doc, "Regular Dividend — the company pays a consistent, predictable dividend every year. Enter the annual dividend per share.", 1)
bullet(doc, "Irregular Dividend — the company pays dividends but they vary year to year. Enter the expected dividend for each specific year.", 1)
bullet(doc, "No Dividend — the company retains all earnings and pays no dividend. In this case, the full EPS is retained and added to book value each year.", 1)
body(doc, "Why does dividend policy matter here? Because dividends reduce the book value of equity. When a company pays out cash as dividends, the equity on the balance sheet shrinks. This affects how quickly book value grows and therefore how much residual income is generated in future years.", indent=True)

step(doc, 6, "Terminal Growth Rate",
     "This is the rate at which you expect the company's residual income to grow beyond your forecast period — forever. Use a conservative number. For a bank in a developing market like Ghana, 3 to 5 percent is reasonable. For a mature bank in the US or UK, analysts use 1 to 3 percent. This rate must always be lower than the Cost of Equity.")

divider(doc)
heading2(doc, "SHEET 4 — Forecast (Watching the Model Build)", color=TEAL)
heading3(doc, "How book value rolls forward each year", color=TEAL)
instructor_note(doc, "Click on the 'Forecast' tab.")
body(doc, "This sheet is the mechanical heart of the model. Watch carefully what happens here.")
body(doc, "In Year 0, we start with the current Book Value Per Share that you entered in Inputs.")
body(doc, "In Year 1, the book value grows as follows:")
formula_box(doc, "Book Value (Year t)  =  Book Value (Year t-1)  +  EPS (Year t)  −  Dividends (Year t)")
body(doc, "Think about what this means. The company earns profit (EPS). It pays some of that out as dividends. Whatever is left — the retained earnings — is added back to the book value. This is exactly how a balance sheet works. Retained earnings accumulate in equity year after year.")
body(doc, "Then, for each year, the model computes the Residual Income:")
formula_box(doc, "RI (Year t)  =  EPS (Year t)  −  Cost of Equity  ×  Book Value (Year t-1)")
body(doc, "And discounts it back to today:")
formula_box(doc, "PV of RI (Year t)  =  RI (Year t)  ÷  (1  +  Cost of Equity)^t")
body(doc, "This is repeated for each year of the forecast period. The model sums all these present values.")
instructor_note(doc, "Trace through Year 1 and Year 2 on screen. Show the formula in the cell. Highlight that book value in Year 2 is different from Year 1 because it grew by retained earnings.")

divider(doc)
heading2(doc, "SHEET 5 — Valuation (The Final Answer)", color=TEAL)
heading3(doc, "Putting it all together", color=TEAL)
instructor_note(doc, "Click on the 'Valuation' tab.")
body(doc, "Now we see the final valuation. The model has three components that add up to give the intrinsic value of the stock:")
formula_box(doc, "Intrinsic Value  =  Current Book Value  +  PV of Residual Income  +  PV of Terminal Value")
body(doc, "Let me walk through each:")
bullet(doc, "Current Book Value — this is what the company is worth today based purely on its net assets. If you liquidated the company right now, this is roughly what you would get per share. This is the floor value.")
bullet(doc, "PV of Residual Income — this is the extra value the company creates above and beyond the shareholders' required return, over your forecast period. If the company consistently earns ROE above the Cost of Equity, this number is positive and adds significant value.")
bullet(doc, "PV of Terminal Value — this captures the residual income generated beyond your forecast period, discounted to today. It is calculated as:")
formula_box(doc, "Terminal Value  =  RI (Final Year)  ×  (1 + g)  ÷  (Cost of Equity  −  g)")
body(doc, "The sum of these three components is your intrinsic value per share. Compare this to the current market price:")
bullet(doc, "If Intrinsic Value > Market Price: stock appears undervalued — potentially a buy")
bullet(doc, "If Intrinsic Value < Market Price: stock appears overvalued — potentially a sell")
bullet(doc, "If Intrinsic Value ≈ Market Price: stock appears fairly priced")
instructor_note(doc, "Show the final intrinsic value. Ask students to compare it to the market price. Have a brief discussion: 'What could cause this stock to be undervalued or overvalued?'")

divider(doc)
heading2(doc, "SHEET 6 — Sensitivity Analysis (The Professional Test)", color=TEAL)
heading3(doc, "Testing two key assumptions simultaneously", color=TEAL)
instructor_note(doc, "Click on the 'Sensitivity' tab.")
body(doc, "Just as in the DCF model, we finish with a sensitivity analysis. But this model has TWO sensitivity tables, not one.")
body(doc, "The first table varies the Long-Term ROE against the Terminal Growth Rate. Why ROE? Because for a bank, long-term return on equity is the single most important driver of value. A bank that consistently earns 20% ROE is far more valuable than one that earns 10% ROE — even if all other things are equal.")
body(doc, "The second table varies Beta against the Terminal Growth Rate. Beta affects the Cost of Equity, which is the discount rate. A higher beta means higher risk, higher required return, lower present value of future cash flows. This table shows how sensitive your valuation is to your estimate of the company's systematic risk.")
body(doc, "Read both tables. Find the cell that corresponds to your base-case assumptions. Now look at the cells around it. This tells you the range of reasonable values for the stock. That range is what you would present to a client, not a single point estimate.")
body(doc, "Remember: every model is wrong, but some models are useful. Your job is not to find the exact right answer. Your job is to build a defensible range of values based on reasonable assumptions — and then explain why you believe the market is wrong or right about where the stock is trading.")
instructor_note(doc, "Point to the 'bull zone' (green cells) and 'bear zone' (red cells). Ask: 'Under what conditions does this stock become a strong buy?'")

divider(doc)
heading2(doc, "Closing — Model 2 Summary and Key Differences", color=TEAL)
body(doc, "Fantastic. Let us summarise Model 2 and contrast it with Model 1.")
bullet(doc, "The Residual Income Model starts from Book Value, not Cash Flow.")
bullet(doc, "It is the preferred model for banks and financial institutions.")
bullet(doc, "The key driver is ROE versus Cost of Equity. If ROE > ke, value is created. If ROE < ke, value is destroyed.")
bullet(doc, "Dividend policy directly affects book value growth and therefore residual income.")
bullet(doc, "The final intrinsic value = Book Value + PV of Residual Income + PV of Terminal Value.")
divider(doc)
heading2(doc, "Final Comparison — Which Model to Use When?")
body(doc, "Before we close, I want to make sure you know when to use each model.")
bullet(doc, "DCF FCFE — use for non-financial companies: manufacturers, retailers, telecoms, consumer goods, energy companies. Any company where you can clearly see free cash flow on the cash flow statement.")
bullet(doc, "Residual Income Model — use for banks, insurance companies, and financial institutions. Their financials are structured differently, and book value and ROE are more meaningful metrics.")
body(doc, "In practice, professional analysts often run BOTH models and triangulate. If both models give you a similar intrinsic value, you have more confidence in the number. If they diverge significantly, you investigate why — are the assumptions different? Is one model more appropriate for this specific company?")
body(doc, "You now have two of the most powerful equity valuation tools used in international capital markets. These are the same models used on Wall Street, Canary Wharf, and every major financial centre in the world. Use them well.")
instructor_note(doc, "End the session here. Open the floor for questions. If time allows, have students pick a company they follow and try entering its numbers live.")

divider(doc)
add_page_break(doc)

# ── Quick Reference Card ───────────────────────────────────────
heading1(doc, "Quick Reference Card — Key Formulas")
heading2(doc, "Model 1: DCF FCFE")
formula_box(doc, "FCFE  =  Net Income  +  D&A  −  Capex  −  ΔNWC  +  Net Borrowings")
formula_box(doc, "Cost of Equity  =  Rf  +  β  ×  (Rm  −  Rf)")
formula_box(doc, "PV of FCFE_t  =  FCFE_t  ÷  (1 + ke)^t")
formula_box(doc, "Terminal Value  =  FCFE_(N+1)  ÷  (ke  −  g)")
formula_box(doc, "Intrinsic Value Per Share  =  (Σ PV of FCFE  +  PV of TV)  ÷  Shares Outstanding")

heading2(doc, "Model 2: Residual Income Model")
formula_box(doc, "Residual Income  =  EPS  −  (ke  ×  Book Value_prev)")
formula_box(doc, "Book Value_t  =  Book Value_(t-1)  +  EPS_t  −  Dividends_t")
formula_box(doc, "PV of RI_t  =  RI_t  ÷  (1 + ke)^t")
formula_box(doc, "Terminal Value  =  RI_final  ×  (1 + g)  ÷  (ke  −  g)")
formula_box(doc, "Intrinsic Value  =  Book Value_0  +  Σ PV of RI  +  PV of TV")

heading2(doc, "Key Rules to Always Remember")
bullet(doc, "Terminal Growth Rate must ALWAYS be less than Cost of Equity — otherwise the model produces an infinite or negative value.")
bullet(doc, "Use sensitivity analysis every time — never present a single-point valuation.")
bullet(doc, "Use DCF FCFE for industrial/commercial companies; use RIM for banks and financial firms.")
bullet(doc, "Beta above 1.0 = more volatile than market. Beta below 1.0 = less volatile than market.")
bullet(doc, "If ROE > Cost of Equity, the company is creating value above and beyond shareholder expectations.")

# ── Save ──────────────────────────────────────────────────────
output_path = r"C:\Users\kkyei\Desktop\Valuation_Lecture_Script.docx"  # noqa
doc.save(output_path)
print(f"Saved: {output_path}")
